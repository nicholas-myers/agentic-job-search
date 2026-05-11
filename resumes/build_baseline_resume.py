#!/usr/bin/env python3
"""Generate baseline job-application resume DOCX. Run from repo root:
   .venv-resume/bin/python resumes/build_baseline_resume.py
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

OUT_PATH = Path(__file__).resolve().parent / "Myers_Baseline_Resume.docx"


def add_center_line(doc: Document, text: str, *, bold: bool = False, size_pt: int = 11) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size_pt)
    p.paragraph_format.space_after = Pt(2)


def add_section_heading(doc: Document, title: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(title.upper())
    run.bold = True
    run.font.size = Pt(11)
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(6)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        doc.add_paragraph(item, style="List Bullet")


def build() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Pt(54)
    section.bottom_margin = Pt(54)
    section.left_margin = Pt(54)
    section.right_margin = Pt(54)

    add_center_line(doc, "Nick Myers", bold=True, size_pt=20)
    add_center_line(doc, "Full-Stack Software Engineer · Portland, OR · Open to remote / hybrid", size_pt=11)
    add_center_line(doc, "https://www.linkedin.com/in/nicholas-myers-professional/", size_pt=10)

    doc.add_paragraph(
        "Full-stack engineer with experience shipping web applications, APIs, integrations, and data "
        "migrations for consulting clients across education, public sector, and commercial programs. "
        "Comfortable owning features end-to-end in React/TypeScript and .NET, partnering with stakeholders "
        "on requirements, and keeping releases testable and maintainable. Previously Senior Programmer/Analyst "
        "at Resource Data; actively seeking software engineering roles."
    )

    add_section_heading(doc, "Technical skills")
    doc.add_paragraph(
        "Languages & frameworks: JavaScript, TypeScript, Python, C#, Java, HTML/CSS, React, Redux, "
        "Node.js, .NET / .NET Core (MVC), Java Spring, jQuery, Bootstrap, SASS/LESS."
    )
    doc.add_paragraph(
        "Data & APIs: SQL Server, GraphQL, REST, JSON, Postman, Boomi, SSMS, data modeling & migrations."
    )
    doc.add_paragraph(
        "Tools & delivery: Git, GitHub, Visual Studio, Agile/Scrum, Storybook, Heroku, dotCMS, Esri maps."
    )

    add_section_heading(doc, "Experience")
    p = doc.add_paragraph()
    r = p.add_run("Resource Data, Inc.")
    r.bold = True
    p.add_run(" — ")
    p.add_run("Senior Programmer/Analyst & Programmer/Analyst")
    p2 = doc.add_paragraph()
    p2.add_run("May 2021 – January 2025 · IT consulting / custom software delivery")
    p2.runs[0].italic = True

    add_bullets(
        doc,
        [
            "Delivered full-stack enhancements for public-sector programs (e.g., Washington Student "
            "Achievement Council Career Launch; Washington DNR burn permitting), improving forms, data "
            "views, and map-driven workflows using React, Redux, and .NET.",
            "Built configurable ordering and document experiences for manufacturing clients (IdeaRoom / "
            "American Steel) with React, TypeScript, and Storybook to streamline custom sales flows.",
            "Improved API reliability for e-commerce-related integrations (Bel/Cinch): GraphQL and .NET "
            "services backed by structured Postman test coverage and stock-data analysis.",
            "Executed student information system transitions for Epic Charter School: Powerschool-to-"
            "curricula migrations, Edgenuity integrations, and Node.js/TypeScript automation; optimized "
            "HTTP concurrency and data-access patterns on .NET and SQL Server.",
            "Extended enterprise CMS capabilities (CUI Inc., dotCMS) with backend work that improved "
            "authoring workflows and content delivery.",
            "Contributed to master-data and integration initiatives for a major agribusiness: Boomi-based "
            "integrations, JSON-oriented hub design, and cross-functional requirements for a marketing data platform.",
        ],
    )

    p = doc.add_paragraph()
    r = p.add_run("Freelance Developer — NimblePath")
    r.bold = True
    doc.add_paragraph("2021 – 2023 · SaaS product development", style=None)
    q = doc.paragraphs[-1]
    q.runs[0].italic = True
    add_bullets(
        doc,
        [
            "Translated wireframes into a React/Redux UI and implemented a Java Spring backend deployed on "
            "Heroku for a scalable SaaS-style architecture.",
        ],
    )

    p = doc.add_paragraph()
    r = p.add_run("Team Lead — Lambda School")
    r.bold = True
    doc.add_paragraph("2020 – 2021 · Technical mentorship", style=None)
    q = doc.paragraphs[-1]
    q.runs[0].italic = True
    add_bullets(
        doc,
        [
            "Mentored ~10 students with weekly 1:1 code reviews, daily standups, and structured feedback to "
            "accelerate learning outcomes.",
        ],
    )

    p = doc.add_paragraph()
    r = p.add_run("Earlier roles")
    r.bold = True
    add_bullets(
        doc,
        [
            "Cable Technician, Spectra Broadband (2019 – 2020): residential installs and troubleshooting; "
            "strong customer satisfaction track record.",
            "Customer Support Representative, Xerox / Verizon program (2016 – 2019): technical phone "
            "support with consistent quality-assurance performance.",
        ],
    )

    add_section_heading(doc, "Certifications")
    add_bullets(
        doc,
        [
            "Boomi Associate MDH; Boomi Associate Developer; Boomi Professional Developer",
            "Full-Stack Web Development and Technical Interviewing — Lambda School",
        ],
    )

    return doc


def main() -> None:
    doc = build()
    doc.save(OUT_PATH)
    print(f"Wrote {OUT_PATH}")


if __name__ == "__main__":
    main()
